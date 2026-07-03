#!/usr/bin/env python3
"""Write the Experimental Results section to a Word document.

The script is intentionally data-driven: it reads pre-generated paper assets
from outputs/paper_results and does not synthesize missing numerical results.
"""
import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SECTIONS = [
    (
        "A. 仿真设置与评估协议",
        ["table_simulation_setup", "table_methods"],
        [],
        "实验在固定轨迹 RGB-D 重建任务下进行，策略仅控制通信资源分配，而不改变无人机飞行轨迹、"
        "视角序列或三维重建流程。因此，实验结果主要反映不同资源分配策略在 RGB 语义传输质量、"
        "深度辅助链路完整度和三维重建质量之间的权衡能力。",
    ),
    (
        "B. GSFusion 代理模型验证",
        "table_surrogate_validation",
        [
            "fig_surrogate_real_vs_pred.png",
            "fig_surrogate_model_comparison.png",
        ],
        "由于在强化学习训练中直接调用完整 GSFusion 重建代价过高，本文采用由真实 GSFusion "
        "结果标定的代理模型。该代理模型只使用 Q_rgb 和 R_depth 作为输入；K_D、beta_D、功率、"
        "tx_amount_norm 等动作相关变量不直接作为代理模型输入，从而避免奖励泄漏。",
    ),
    (
        "C. 基于验证集的超参数选择",
        "table_validation_selection",
        ["fig_ppo_penalty_validation_sensitivity.png"],
        "PPO-penalty 的惩罚权重在验证轨迹上选择，而不是在测试轨迹上调参。该流程避免了测试集泄漏，"
        "并为 PPO-penalty 提供公平的对比设置。",
    ),
    (
        "D. 未见轨迹上的主要测试结果",
        "table_main_test_performance",
        [
            "fig_main_q3d_comparison.png",
            "fig_main_constraint_violation.png",
            "fig_quality_constraint_tradeoff.png",
        ],
        None,
    ),
    (
        "E. 约束行为与资源分配分析",
        None,
        [
            "fig_resource_allocation_comparison.png",
            "fig_crpo_mode_ratio.png",
            "fig_crpo_training_curves.png",
            "fig_lagrangian_lambda_evolution.png",
        ],
        "CRPO-PPO 在重建奖励优化和约束修复之间切换：当 RGB 或深度约束被违反时，更新目标转向"
        "降低最严重违反的服务质量代价；当约束满足时，更新目标回到提升 q3D。这里 K_D 表示分配给"
        "数字深度链路的子载波数，beta_D 表示分配给深度链路的功率占比。",
    ),
    (
        "F. 敏感性与消融实验",
        ["table_sensitivity", "table_crpo_epsilon_ablation"],
        [
            "fig_penalty_lambda_sensitivity.png",
            "fig_crpo_epsilon_sensitivity.png",
        ],
        None,
    ),
]


TABLE_CAPTIONS = {
    "table_simulation_setup": "表 I  仿真设置与评估协议。",
    "table_methods": "表 II  对比方法设置。",
    "table_surrogate_validation": "表 III  GSFusion 代理模型验证结果。",
    "table_validation_selection": "表 IV  基于验证集的 PPO-penalty 超参数选择。",
    "table_main_test_performance": "表 V  未见测试轨迹上的主要性能。",
    "table_sensitivity": "表 VI  PPO-penalty 惩罚权重敏感性。",
    "table_crpo_epsilon_ablation": "表 VII  CRPO-PPO epsilon 敏感性验证。",
}


FIGURE_CAPTIONS = {
    "fig_surrogate_real_vs_pred.png": "图 1  真实 GSFusion 重建质量与代理模型预测结果对比。",
    "fig_surrogate_model_comparison.png": "图 2  线性代理模型与饱和代理模型的性能对比。",
    "fig_ppo_penalty_validation_sensitivity.png": "图 3  PPO-penalty 对惩罚权重的验证集敏感性。",
    "fig_main_q3d_comparison.png": "图 4  不同方法的三维重建质量对比。",
    "fig_main_constraint_violation.png": "图 5  未见测试轨迹上的 RGB 与深度约束代价。",
    "fig_quality_constraint_tradeoff.png": "图 6  三维重建质量与约束违反之间的权衡关系。",
    "fig_resource_allocation_comparison.png": "图 7  不同方法的子载波与功率分配行为。",
    "fig_crpo_mode_ratio.png": "图 8  CRPO-PPO 训练过程中的模式分布。",
    "fig_crpo_training_curves.png": "图 9  CRPO-PPO 的奖励与约束代价训练曲线。",
    "fig_lagrangian_lambda_evolution.png": "图 10  Lagrangian-PPO 对偶变量演化。",
    "fig_penalty_lambda_sensitivity.png": "图 11  PPO-penalty 的 lambda 敏感性分析。",
    "fig_crpo_epsilon_sensitivity.png": "图 12  CRPO-PPO 的 epsilon 敏感性分析。",
}


def clean_inline(text):
    text = text.replace("~\\ref", " ")
    text = re.sub(r"\\ref\{[^}]+\}", "", text)
    text = re.sub(r"Table\s*~", "Table ", text)
    text = re.sub(r"Fig\.\s*~", "Fig. ", text)
    text = text.replace("−", "-")
    return " ".join(text.split())


def split_analysis_sections(markdown_text):
    sections = {}
    current = None
    buf = []
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if line.startswith("### "):
            if current:
                sections[current] = "\n".join(buf).strip()
            current = normalize_section_title(line[4:].strip())
            buf = []
        elif current:
            buf.append(raw_line.rstrip())
    if current:
        sections[current] = "\n".join(buf).strip()
    return sections


def normalize_section_title(title):
    title = title.strip()
    if title.startswith("A."):
        return "A"
    if title.startswith("B."):
        return "B"
    if title.startswith("C."):
        return "C"
    if title.startswith("D."):
        return "D"
    if title.startswith("E."):
        return "E"
    if title.startswith("F."):
        return "F"
    return title


def add_paragraphs_from_markdown(doc, markdown_chunk):
    paragraphs = []
    buf = []
    for raw_line in markdown_chunk.splitlines():
        line = raw_line.strip()
        if not line:
            if buf:
                paragraphs.append(" ".join(buf))
                buf = []
            continue
        if line.startswith("**") and line.endswith("**"):
            if buf:
                paragraphs.append(" ".join(buf))
                buf = []
            paragraphs.append(line.strip("*"))
            continue
        buf.append(line)
    if buf:
        paragraphs.append(" ".join(buf))

    for para in paragraphs:
        text = clean_inline(para)
        if not text:
            continue
        if len(text) < 80 and text.endswith(".") is False and ":" not in text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        else:
            doc.add_paragraph(text)


def set_table_style(table):
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def add_csv_table(doc, table_path):
    with open(table_path, newline="") as f:
        rows = list(csv.reader(f))
    if not rows:
        doc.add_paragraph("This result is not available in the current outputs.")
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_style(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def find_table(paper_results, stem):
    for suffix in (".csv", ".md", ".tex"):
        candidate = paper_results / "tables" / f"{stem}{suffix}"
        if candidate.exists():
            return candidate
    return None


def add_table(doc, paper_results, stem, inserted, missing):
    if not stem:
        return
    caption = TABLE_CAPTIONS.get(stem, f"Table. {stem}.")
    table_path = find_table(paper_results, stem)
    if table_path is None:
        doc.add_paragraph(caption)
        doc.add_paragraph("This result is not available in the current outputs.")
        missing.append(stem)
        return
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.bold = True
    if table_path.suffix == ".csv":
        add_csv_table(doc, table_path)
    else:
        doc.add_paragraph(table_path.read_text(encoding="utf-8"))
    inserted.append(str(table_path.relative_to(paper_results)))


def add_figure(doc, paper_results, figure_name, inserted, missing):
    figure_path = paper_results / "figures" / figure_name
    if not figure_path.exists():
        missing.append(figure_name)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(figure_path), width=Inches(6.0))
    cap = doc.add_paragraph(FIGURE_CAPTIONS.get(figure_name, f"Fig. {figure_name}."))
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
    inserted.append(str(figure_path.relative_to(paper_results)))


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Times New Roman"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 2"].font.name = "Times New Roman"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Heading 2"].font.size = Pt(12)


def write_cn_summary(path, inserted_tables, inserted_figures, missing_tables, missing_figures):
    lines = [
        "# Experimental Results Word 生成说明",
        "",
        "## 已生成小节",
        "- V. 实验结果",
        "- A. 仿真设置与评估协议",
        "- B. GSFusion 代理模型验证",
        "- C. 基于验证集的超参数选择",
        "- D. 未见轨迹上的主要测试结果",
        "- E. 约束行为与资源分配分析",
        "- F. 敏感性与消融实验",
        "",
        "## 已插入表格",
    ]
    lines += [f"- {x}" for x in inserted_tables] or ["- 无"]
    lines += ["", "## 已插入图片"]
    lines += [f"- {x}" for x in inserted_figures] or ["- 无"]
    lines += ["", "## 缺失表格"]
    lines += [f"- {x}" for x in missing_tables] or ["- 无"]
    lines += ["", "## 缺失图片"]
    lines += [f"- {x}" for x in missing_figures] or ["- 无"]
    lines += [
        "",
        "## 当前结果解释建议",
        "- 本结果部分应表述为 constrained DRL 方法的系统比较，不应强行声称 CRPO-PPO 全面优于 PPO-Penalty。",
        "- PPO-Penalty 的结果应强调经过 validation-based penalty tuning 后得到。",
        "- CRPO-PPO 的重点是显式约束处理和可解释的 constraint-rectification 行为。",
        "- surrogate validation 部分应强调代理模型只使用 Q_rgb 和 R_depth，未直接使用 K_D、beta_D、power 或 tx_amount_norm，从而避免 reward leakage。",
        "- 缺失的 sensitivity/ablation 图表应标记为 TODO，不应编造实验数据。",
    ]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--paper-results", default="outputs/paper_results")
    parser.add_argument("--out", default="outputs/experimental_results_section.docx")
    args = parser.parse_args()

    paper_results = Path(args.paper_results)
    if not paper_results.exists():
        raise SystemExit(
            "outputs/paper_results/ does not exist. Please run:\n"
            "python scripts/generate_paper_results.py \\\n"
            "  --results-root outputs \\\n"
            "  --config configs/default.yaml \\\n"
            "  --out-dir outputs/paper_results"
        )

    analysis_path = paper_results / "analysis_cn.md"
    if not analysis_path.exists():
        analysis_path = paper_results / "analysis_en.md"
    if not analysis_path.exists():
        raise SystemExit(f"Missing required analysis file: {analysis_path}")

    markdown = analysis_path.read_text(encoding="utf-8")
    analysis_sections = split_analysis_sections(markdown)

    doc = Document()
    configure_document(doc)
    doc.add_heading("V. 实验结果", level=1)

    inserted_tables = []
    missing_tables = []
    inserted_figures = []
    missing_figures = []

    for title, table_stems, figures, required_text in SECTIONS:
        doc.add_heading(title, level=2)
        if required_text:
            doc.add_paragraph(clean_inline(required_text))
        chunk = analysis_sections.get(normalize_section_title(title), "")
        if chunk:
            add_paragraphs_from_markdown(doc, chunk)
        else:
            doc.add_paragraph("当前输出中没有该实验结果。")
        if table_stems is None:
            table_stems = []
        elif isinstance(table_stems, str):
            table_stems = [table_stems]
        for table_stem in table_stems:
            add_table(doc, paper_results, table_stem, inserted_tables, missing_tables)
        for figure_name in figures:
            add_figure(doc, paper_results, figure_name, inserted_figures, missing_figures)
        if title.startswith("F.") and "fig_crpo_epsilon_sensitivity.png" in missing_figures:
            doc.add_paragraph(
                "当前输出中没有对应的敏感性结果，后续实验应补充生成。"
            )

    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    summary_path = out_path.parent / "experimental_results_section_cn_summary.md"
    write_cn_summary(summary_path, inserted_tables, inserted_figures, missing_tables, missing_figures)

    # Also write a machine-readable summary next to the Word file for auditing.
    audit = {
        "inserted_tables": inserted_tables,
        "inserted_figures": inserted_figures,
        "missing_tables": missing_tables,
        "missing_figures": missing_figures,
    }
    audit_path = out_path.with_suffix(".summary.json")
    audit_path.write_text(json.dumps(audit, indent=2), encoding="utf-8")

    print(f"Generated Word file: {out_path}")
    print(f"Generated Chinese summary: {summary_path}")
    print("Inserted tables:", ", ".join(inserted_tables) if inserted_tables else "none")
    print("Inserted figures:", ", ".join(inserted_figures) if inserted_figures else "none")
    print("Missing tables:", ", ".join(missing_tables) if missing_tables else "none")
    print("Missing figures:", ", ".join(missing_figures) if missing_figures else "none")


if __name__ == "__main__":
    main()
